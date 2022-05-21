from docx import Document

document = Document()

document.add_heading('簡単なWordドキュメントのタイトル', 0)
document.add_paragraph('簡単なWordドキュメントのテキスト')

document.add_picture('png.png')

count = 0
for data in document.paragraphs:
    count += len(data.text)
print(f'sample.docxの中の文字数は{count}個です。')

document.save('sample.docx')